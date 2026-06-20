# -*- coding: utf-8 -*-
"""보고서_본문.md → docx 변환. 템플릿(final_exam보고서 양식.docx) 서식 재현.
구문: # H1 / ## H2 / ### H3 / [[FIG path|caption]] / [[SYS ...]] / 파이프표 / - 불릿 / **bold** / --- 무시.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
MD = ROOT/"보고서_본문.md"
OUT = ROOT/"판교_청라_업무지구_비교분석_보고서.docx"
FONT = "맑은 고딕"
BODY_PT = 11
INK = RGBColor(0x1a,0x1a,0x1a)
ACCENT = RGBColor(0x1d,0x4e,0xd8)
GRAY = RGBColor(0x55,0x5b,0x66)

def set_run(run, size=BODY_PT, bold=False, italic=False, color=INK, name=FONT):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = color; run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rf.set(qn(a), name)

def add_runs(p, text, size=BODY_PT, color=INK, base_bold=False):
    """**bold** 인라인 처리."""
    for i, seg in enumerate(re.split(r'\*\*', text)):
        if seg == "": continue
        r = p.add_run(seg)
        set_run(r, size=size, bold=(base_bold or i % 2 == 1), color=color)

def shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),fill)
    pPr.append(shd)

def left_border(p, color="1d4ed8", sz="18"):
    pPr = p._p.get_or_add_pPr()
    pbd = OxmlElement('w:pBdr'); left = OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),sz); left.set(qn('w:space'),'8'); left.set(qn('w:color'),color)
    pbd.append(left); pPr.append(pbd)

doc = Document()
# --- 페이지: A4 + 1인치 여백 ---
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
for m in ('left_margin','right_margin','top_margin','bottom_margin'):
    setattr(sec, m, Cm(2.4))
# Normal 기본 폰트
nstyle = doc.styles['Normal']; nstyle.font.name = FONT; nstyle.font.size = Pt(BODY_PT)
nstyle.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def para(align=None, before=0, after=6, line=1.42):
    p = doc.add_paragraph(); pf = p.paragraph_format
    if align is not None: p.alignment = align
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    return p

# ============ 표지 ============
def cover():
    for _ in range(3): doc.add_paragraph()
    p = para(WD_ALIGN_PARAGRAPH.CENTER, after=4); add_runs(p, "수업 프로젝트 보고서", size=13, color=GRAY)
    p = para(WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4)
    r = p.add_run("데이터로 진단하는 업무지구의 성공과 실패"); set_run(r, size=23, bold=True, color=INK)
    p = para(WD_ALIGN_PARAGRAPH.CENTER, after=2)
    r = p.add_run("— 판교테크노밸리와 청라국제도시의 정량 비교분석 —"); set_run(r, size=13, color=GRAY)
    p = para(WD_ALIGN_PARAGRAPH.CENTER, before=6, after=2)
    r = p.add_run("핵심역의 노동시장 접근성을 출발점으로 한 성공요인 도출"); set_run(r, size=11, italic=True, color=GRAY)
    for _ in range(4): doc.add_paragraph()
    # 메타 표
    meta = [("과 목 명","스마트시티 이론과 실제 (여지호 교수)"),
            ("분반 / 팀(조)","개인 과제 (Take-Home)"),
            ("팀원 (학번 · 성명)","202435333 · 김태우"),
            ("지도교수","여지호"),
            ("제 출 일","2026년 6월 22일")]
    t = doc.add_table(rows=len(meta), cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    t.columns[0].width = Cm(4.6); t.columns[1].width = Cm(9.4)
    for i,(k,v) in enumerate(meta):
        c0, c1 = t.rows[i].cells
        c0.width = Cm(4.6); c1.width = Cm(9.4)
        p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(k); set_run(r0, size=10.5, bold=True, color=INK)
        shade(p0, "F1F5F9")
        p1 = c1.paragraphs[0]; r1 = p1.add_run(v); set_run(r1, size=10.5, color=INK)
    doc.add_page_break()

# ============ 목차 ============
def toc(headings):
    p = para(WD_ALIGN_PARAGRAPH.CENTER, after=12); r = p.add_run("목   차"); set_run(r, size=16, bold=True, color=INK)
    for level, text in headings:
        pp = para(after=4, line=1.3)
        pp.paragraph_format.left_indent = Cm(0.0 if level==1 else 0.8)
        add_runs(pp, text, size=11.5 if level==1 else 10.5, color=INK, base_bold=(level==1))
    doc.add_page_break()

# ============ 그림 ============
FIG_W = Cm(10.8)
SYSFIG_W = Cm(15.4)
def figure(path, caption, width=None):
    p = para(WD_ALIGN_PARAGRAPH.CENTER, before=3, after=1)
    p.add_run().add_picture(str((ROOT/path).resolve()), width=width or FIG_W)
    cap = para(WD_ALIGN_PARAGRAPH.CENTER, after=5, line=1.1)
    add_runs(cap, caption, size=8.6, color=GRAY)

# ============ 표 ============
def cell_pad(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for side in ('top','bottom'):
        e = OxmlElement(f'w:{side}'); e.set(qn('w:w'),'20'); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def table(rows):
    header, data = rows[0], rows[1:]
    t = doc.add_table(rows=len(data)+1, cols=len(header)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(header):
        cell = t.rows[0].cells[j]; cell_pad(cell); p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
        add_runs(p, h, size=9.6, color=INK, base_bold=True); shade(p, "E8EDF5")
    for i,row in enumerate(data):
        for j,val in enumerate(row):
            cell = t.rows[i+1].cells[j]; cell_pad(cell); p = cell.paragraphs[0]
            p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (j>0 or len(header)<=2) else WD_ALIGN_PARAGRAPH.LEFT
            add_runs(p, val, size=9.4, color=INK)
    sp = para(after=4, line=1.0)

def heading(text, level):
    sizes = {1:15, 2:12.5, 3:11.5}
    p = para(before=(10 if level==1 else 7), after=(5 if level==1 else 3), line=1.15)
    r = p.add_run(text); set_run(r, size=sizes[level], bold=True, color=(ACCENT if level<=2 else INK))
    if level==1:
        pPr=p._p.get_or_add_pPr(); pbd=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:space'),'4'); bot.set(qn('w:color'),'c7d2e8')
        pbd.append(bot); pPr.append(pbd)

def callout(text):
    p = para(before=3, after=6, line=1.2); p.paragraph_format.left_indent=Cm(0.25)
    shade(p, "EEF3FB"); left_border(p)
    add_runs(p, text, size=9.7, color=RGBColor(0x33,0x3a,0x46))

# ============ 파싱 ============
raw = MD.read_text(encoding="utf-8").split("\n")
headings = []
for ln in raw:
    m = re.match(r'^(#{1,3})\s+(.*)$', ln.strip())
    if m: headings.append((len(m.group(1)), m.group(2).strip()))

cover(); toc(headings)

i = 0; n = len(raw)
while i < n:
    line = raw[i].rstrip()
    s = line.strip()
    if s == "" or s == "---":
        i += 1; continue
    # 헤딩
    m = re.match(r'^(#{1,3})\s+(.*)$', s)
    if m:
        heading(m.group(2).strip(), len(m.group(1))); i += 1; continue
    # 그림
    mf = re.match(r'^\[\[FIG\s+(.+?)\s*\|\s*(.+?)\]\]$', s)
    if mf:
        figure(mf.group(1).strip(), mf.group(2).strip()); i += 1; continue
    # 시스템 화면(와이드)
    ms = re.match(r'^\[\[SYSFIG\s+(.+?)\s*\|\s*(.+?)\]\]$', s)
    if ms:
        figure(ms.group(1).strip(), ms.group(2).strip(), width=SYSFIG_W); i += 1; continue
    # 시스템 연계 (blockquote)
    if s.startswith(">"):
        callout(re.sub(r'^>\s*', '', s)); i += 1; continue
    msys = re.match(r'^\[\[SYS\s+(.+?)\]\]$', s)
    if msys:
        callout(msys.group(1).strip()); i += 1; continue
    # 표
    if s.startswith("|"):
        block = []
        while i < n and raw[i].strip().startswith("|"):
            block.append(raw[i].strip()); i += 1
        rows = []
        for b in block:
            if re.match(r'^\|[\s:\-\|]+\|$', b): continue   # separator
            cells = [c.strip() for c in b.strip().strip('|').split('|')]
            rows.append(cells)
        if rows: table(rows)
        continue
    # 불릿
    if s.startswith("- "):
        p = para(WD_ALIGN_PARAGRAPH.JUSTIFY, after=2, line=1.26); p.paragraph_format.left_indent=Cm(0.5)
        p.paragraph_format.first_line_indent=Cm(-0.25)
        rb=p.add_run("•  "); set_run(rb, size=BODY_PT, color=ACCENT)
        add_runs(p, s[2:], size=BODY_PT, color=INK); i += 1; continue
    # 일반 문단
    p = para(WD_ALIGN_PARAGRAPH.JUSTIFY, after=3, line=1.27)
    add_runs(p, s, size=BODY_PT, color=INK); i += 1

doc.save(OUT)
print("saved:", OUT)
print("headings:", len(headings), "| paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
